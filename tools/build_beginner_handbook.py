"""Build the OpenRobot-One Day 5-30 beginner handbook."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "OpenRobot-One_Day5-Day30_零基础实操手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
NOTE_FILL = "F4F6F9"
WARN_FILL = "FFF4CE"
RISK_FILL = "FDE7E9"
GREEN_FILL = "E6F4EA"
WHITE = "FFFFFF"
GRID = "B8C2CC"
BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
STEP_COUNTER = 0


def set_run_font(run, name=BODY_FONT, size=11, bold=None, italic=None, color=None):
    """Apply cross-renderer run fonts."""
    run.font.name = name
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        cell._tc.get_or_add_tcPr().append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, color=GRID, size=6):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), color)
        borders.append(edge)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for item in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(item)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.text = "OpenRobot-One｜Day 5–30 零基础实操手册"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.runs[0], size=9, bold=True, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    prefix = fp.add_run("示例参数必须以实测结果为准    ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_number(fp)


def para(doc, text="", bold_prefix=None, align=None, size=11, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=size, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color)
    return p


def heading(doc, text, level=1):
    global STEP_COUNTER
    STEP_COUNTER = 0
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}[level],
            bold=True,
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        )
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_run_font(p.add_run(text), size=11)
    return p


def numbered(doc, text):
    global STEP_COUNTER
    STEP_COUNTER += 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.19)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(
        p.add_run(f"步骤 {STEP_COUNTER}｜"),
        size=10.5,
        bold=True,
        color=DARK_BLUE,
    )
    set_run_font(p.add_run(text), size=11)
    return p


def checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.19)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run("□ "), size=11, bold=True, color=BLUE)
        set_run_font(p.add_run(item), size=11)


def callout(doc, label, text, kind="note"):
    colors = {
        "note": (NOTE_FILL, DARK_BLUE),
        "warning": (WARN_FILL, "7A5A00"),
        "risk": (RISK_FILL, "9B1C1C"),
        "success": (GREEN_FILL, "1F5E32"),
    }
    fill, ink = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color=ink, size=5)
    set_run_font(p.add_run(f"{label}："), size=10.5, bold=True, color=ink)
    set_run_font(p.add_run(text), size=10.5, color=NAVY)
    return p


def command(doc, terminal, cwd, text, expected=None):
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(5)
    label.paragraph_format.space_after = Pt(2)
    label.paragraph_format.keep_with_next = True
    set_run_font(label.add_run(f"终端：{terminal}"), size=9.5, bold=True, color=DARK_BLUE)
    if cwd:
        set_run_font(label.add_run(f"｜当前目录：{cwd}"), size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, "F7F8FA")
    set_paragraph_border(p, color="CCD3DA", size=4)
    for index, line in enumerate(text.strip().splitlines()):
        if index:
            p.add_run().add_break()
        set_run_font(p.add_run(line), name=CODE_FONT, size=9, color="202124")
    if expected:
        callout(doc, "预期结果", expected, "success")


def code(doc, title, text):
    label = para(doc, title, size=9.5, color=DARK_BLUE)
    label.paragraph_format.space_before = Pt(5)
    label.paragraph_format.space_after = Pt(2)
    label.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    # Short snippets read best as one block. Long XML/YAML examples may span
    # pages; forcing those together creates a large blank area on the previous
    # page.
    p.paragraph_format.keep_together = len(text.strip().splitlines()) <= 18
    set_paragraph_shading(p, "F7F8FA")
    set_paragraph_border(p, color="CCD3DA", size=4)
    for index, line in enumerate(text.strip().splitlines()):
        if index:
            p.add_run().add_break()
        set_run_font(p.add_run(line), name=CODE_FONT, size=8.3, color="202124")


def table(doc, headers, rows, widths, first_col_align=WD_ALIGN_PARAGRAPH.CENTER):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        set_run_font(p.add_run(text), size=9.5, bold=True, color=NAVY)
    set_repeat_table_header(tbl.rows[0])
    for row_data in rows:
        cells = tbl.add_row().cells
        for index, text in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index else first_col_align
            set_run_font(p.add_run(str(text)), size=9.2)
    set_table_geometry(tbl, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def page_break(doc):
    doc.add_page_break()


def day_start(doc, day, title, goal):
    heading(doc, f"Day {day}｜{title}", 2)
    callout(doc, "本日完成标准", goal, "success")


def add_cover(doc):
    for _ in range(4):
        para(doc, "")
    p = para(doc, "OPENROBOT-ONE", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=BLUE)
    p.paragraph_format.space_after = Pt(18)
    title = para(
        doc,
        "Day 5–30 零基础实操手册",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=28,
        color=NAVY,
    )
    title.runs[0].bold = True
    title.paragraph_format.space_after = Pt(10)
    subtitle = para(
        doc,
        "从启动开发环境，到仿真、STM32、串口驱动、SLAM 与 Nav2",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        color=DARK_BLUE,
    )
    subtitle.paragraph_format.space_after = Pt(36)
    para(doc, "适用环境", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=MUTED)
    para(
        doc,
        "Windows 11｜WSL2 Ubuntu 22.04｜ROS 2 Humble｜Gazebo Classic 11",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
        color=NAVY,
    )
    para(
        doc,
        "C++17｜Python 3｜Docker｜STM32F407",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
        color=NAVY,
    )
    para(doc, "")
    callout(
        doc,
        "阅读承诺",
        "本手册假定你实操基础薄弱。每条关键命令都会说明在哪个终端运行、"
        "当前目录、预期结果和失败后的恢复办法。后续功能是你要亲手实现的目标，"
        "不是当前仓库已经完成的能力。",
        "note",
    )
    para(doc, "版本：2026-07-29", align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, color=MUTED)
    page_break(doc)


def add_front_matter(doc):
    heading(doc, "开始之前", 1)
    callout(
        doc,
        "当前边界",
        "仓库当前完成 Day 1–4：工程基座、参数化模型、Gazebo 差速仿真和统一 "
        "Bringup。没有 /scan、SLAM、AMCL、Nav2、STM32 固件或真机串口驱动。",
        "warning",
    )
    heading(doc, "你会用到的五类终端", 2)
    table(
        doc,
        ["标记", "在哪里打开", "主要用途"],
        [
            ("P", "Windows PowerShell", "安装/检查 WSL、USBIPD、Windows 软件"),
            ("W", "Ubuntu-22.04 / WSL2", "进入仓库、调用 Docker、挂载 USB"),
            ("C", "Docker 容器 Bash", "执行 ROS 2、colcon、Gazebo、RViz、Nav2"),
            ("S", "STM32CubeIDE", "生成、编译、烧录和调试 STM32 固件"),
            ("R", "串口终端/逻辑分析工具", "观察串口字节、波形和电机反馈"),
        ],
        [1100, 2800, 5460],
    )
    callout(
        doc,
        "最常见的新手错误",
        "把 PowerShell、WSL 和容器命令混在一个窗口执行。看到本手册的“终端”标签后，"
        "先确认窗口类型，再输入命令。",
        "risk",
    )
    heading(doc, "学习路线与阶段门槛", 2)
    table(
        doc,
        ["阶段", "天数", "必须先通过的门槛"],
        [
            ("复现", "Day 1–4", "7 个包构建；59 项测试；/odom 和 TF 正常"),
            ("仿真导航基础", "Day 5–7", "/scan 10 Hz；保存可重复加载的地图"),
            ("固件", "Day 8–14", "双电机闭环；协议稳定；500 ms 超时停车"),
            ("真机 ROS", "Day 15–21", "/cmd_vel 到轮速；/odom、TF、重连正常"),
            ("导航交付", "Day 22–30", "10 条路线统计；新环境复现；发布材料完整"),
        ],
        [1500, 1400, 6460],
    )
    heading(doc, "目录", 2)
    for item in (
        "第 1 章：环境从零搭建与每天如何启动",
        "第 2 章：Day 1–4 到底做了什么，以及如何亲自复现",
        "第 3 章：Day 5–7 激光雷达、办公室世界与 SLAM",
        "第 4 章：Day 8–14 STM32 电机闭环",
        "第 5 章：Day 15–21 ROS 2 串口驱动与里程计",
        "第 6 章：Day 22–30 AMCL、Nav2、测试与发布",
        "第 7 章：按现象排错",
        "附录：参数记录表、命令速查、Topic/TF 所有权和术语",
    ):
        bullet(doc, item)
    callout(doc, "Word 提示", "打开文档后按 Ctrl+A，再按 F9，可刷新页码等字段。", "note")
    page_break(doc)


def add_environment(doc):
    heading(doc, "第 1 章｜环境从零搭建与每天如何启动", 1)
    heading(doc, "1.1 先理解三层环境", 2)
    para(
        doc,
        "Windows 是宿主机；WSL2 提供 Ubuntu 命令行；Docker 容器提供固定版本的 "
        "ROS 2 Humble。项目源码保存在 D 盘，通过 /mnt/d/OpenRobot-One 被 WSL "
        "看到，再挂载到容器的 /workspace。这样本地和 GitHub Actions 使用同一镜像。",
    )
    table(
        doc,
        ["你看到的路径", "真实位置", "用途"],
        [
            (r"D:\OpenRobot-One", "Windows 文件系统", "用编辑器查看和修改源码"),
            ("/mnt/d/OpenRobot-One", "WSL 对 D 盘的映射", "从 Ubuntu 启动 Docker"),
            ("/workspace", "容器挂载点", "ROS 2 构建、测试和运行的根目录"),
            ("/workspace/ros2_ws/src", "容器内源码目录", "7 个 ROS 2 包"),
        ],
        [2600, 2800, 3960],
    )

    heading(doc, "1.2 安装和检查 WSL2", 2)
    callout(doc, "权限", "下面两条安装命令需要“以管理员身份运行”的 PowerShell。", "warning")
    command(
        doc,
        "P｜管理员 PowerShell",
        r"C:\Windows\System32",
        "wsl --install -d Ubuntu-22.04\nwsl --set-default-version 2",
        "首次安装后按提示重启。之后打开普通 PowerShell，wsl -l -v 应显示 Ubuntu-22.04 且 VERSION 为 2。",
    )
    command(doc, "P｜普通 PowerShell", None, "wsl -l -v\nwsl -d Ubuntu-22.04")
    para(doc, "进入 Ubuntu 后，提示符通常以 `$` 结尾。退出 WSL 输入 `exit`。")

    heading(doc, "1.3 Docker Desktop 与 WSL 集成", 2)
    numbered(doc, "安装 Docker Desktop，设置中启用 Use the WSL 2 based engine。")
    numbered(doc, "Settings → Resources → WSL Integration，打开 Ubuntu-22.04。")
    numbered(doc, "等待左下角状态变为 Engine running，再在 WSL 验证。")
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "cd /mnt/d/OpenRobot-One\ndocker version\ndocker compose version",
        "Client 和 Server 都有版本号；不能只有 Client。",
    )
    callout(
        doc,
        "若提示 cannot connect",
        "先确认 Docker Desktop 已启动，再执行 wsl --shutdown，重新打开 Ubuntu；不要使用 sudo 安装第二套 Docker。",
        "risk",
    )

    heading(doc, "1.4 获取并检查项目", 2)
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "cd /mnt/d/OpenRobot-One\ngit status --short --branch\nfind ros2_ws/src -maxdepth 2 -name package.xml -print",
        "分支应为 main；可以看到 7 个 openrobot_* 包。当前仓库可能有未跟踪文件，先不要删除或覆盖。",
    )

    heading(doc, "1.5 构建开发镜像", 2)
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "docker build -f docker/Dockerfile -t openrobot-one:humble .",
        "最后出现 Successfully tagged openrobot-one:humble 或等价成功信息。",
    )
    para(doc, "第一次构建需要下载镜像和 apt 包，时间较长；之后会使用缓存。")

    heading(doc, "1.6 两种进入容器的方法", 2)
    para(doc, "方法 A 适合日常交互开发：")
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "docker compose -f docker/compose.yaml run --rm dev",
        "提示符进入容器，pwd 输出 /workspace，ros2 --help 可用。",
    )
    para(doc, "方法 B 适合一次性构建测试：")
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        'docker run --rm -v "$(pwd):/workspace" openrobot-one:humble \\\n'
        '  bash -lc "./scripts/build_ros.sh"',
        "7 个包构建完成，最终显示 59 tests, 0 errors, 0 failures。",
    )
    callout(
        doc,
        "容器为何能找到 ROS",
        "docker/entrypoint.sh 会先 source /opt/ros/humble/setup.bash；若 /workspace/install/setup.bash "
        "存在，还会加载工作区 overlay。",
        "note",
    )

    heading(doc, "1.7 每天开始开发的固定动作", 2)
    checklist(
        doc,
        [
            "启动 Docker Desktop，等到 Engine running。",
            "打开 Ubuntu-22.04，执行 cd /mnt/d/OpenRobot-One。",
            "执行 git status --short --branch，确认没有误覆盖自己的修改。",
            "进入 dev 容器，执行 pwd 和 ros2 pkg list | grep openrobot。",
            "修改前先打开相关 README、package.xml、CMakeLists.txt、Launch、Xacro 和测试。",
            "先写成功标准，再做最小修改。",
        ],
    )
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "source /opt/ros/humble/setup.bash\n"
        "[ -f install/setup.bash ] && source install/setup.bash\n"
        "colcon list --base-paths ros2_ws/src",
        "列出 openrobot_bringup、description、driver、gazebo、msgs、navigation、tests。",
    )

    heading(doc, "1.8 构建、测试、启动、停止", 2)
    command(doc, "C｜容器 Bash", "/workspace", "./scripts/build_ros.sh")
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 launch openrobot_bringup bringup.launch.py \\\n"
        "  sim:=true use_sim_time:=true use_rviz:=false",
        "Gazebo 服务启动，spawn_entity 报告机器人生成成功。终端保持占用是正常现象。",
    )
    para(doc, "停止当前程序：回到启动它的终端，按一次 Ctrl+C，等待所有子进程退出。不要直接关闭 Docker Desktop。")
    callout(
        doc,
        "多终端规则",
        "Launch 占用终端 1；Topic/TF 检查使用终端 2；键盘控制使用终端 3。每个新容器终端都要 source install/setup.bash。",
        "warning",
    )
    command(
        doc,
        "W｜新 WSL 终端",
        "/mnt/d/OpenRobot-One",
        "docker ps --format 'table {{.ID}}\\t{{.Image}}\\t{{.Names}}'\n"
        "docker exec -it <容器名> bash\n"
        "source /workspace/install/setup.bash",
    )
    para(doc, "若使用 `docker compose run --rm dev`，先用 `docker ps` 找到自动生成的容器名，再 exec 进入同一个容器，ROS DDS 才最省事。")

    heading(doc, "1.9 GUI、RViz 和 Gazebo 客户端", 2)
    para(
        doc,
        "当前最可靠的学习路径是先用 use_rviz:=false 完成无界面验收。GUI 需要 WSLg/X11 "
        "转发；项目 compose 只设置 DISPLAY，主机配置不同时可能不能直接显示。",
    )
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "echo \"$DISPLAY\"\nls -la /tmp/.X11-unix",
        "DISPLAY 非空且 X11 socket 存在，才继续尝试 GUI。",
    )
    callout(
        doc,
        "GUI 失败不等于 ROS 失败",
        "先验证 /odom、/joint_states 和 TF。若这些正常，只把 RViz/Gazebo 客户端标记为未验证，"
        "不要为显示问题改机器人算法。",
        "note",
    )

    heading(doc, "1.10 STM32 与 USB 工具准备", 2)
    bullet(doc, "Windows 安装 STM32CubeIDE 和 ST-Link 驱动；先用空工程确认能识别 ST-Link。")
    bullet(doc, "Windows 安装 usbipd-win。只有需要让 WSL/容器直接访问 USB 时才执行 USB 转发。")
    bullet(doc, "CubeIDE 原生 Windows 调试时，ST-Link 留在 Windows；不要同时 attach 给 WSL。")
    command(doc, "P｜管理员 PowerShell", None, "usbipd list")
    command(
        doc,
        "P｜管理员 PowerShell",
        None,
        "usbipd bind --busid <BUSID>\nusbipd attach --wsl --busid <BUSID>",
        "WSL 中 lsusb 能看到设备。BUSID 必须来自 usbipd list，不能猜。",
    )
    command(doc, "W｜Ubuntu-22.04", None, "lsusb\nls -l /dev/ttyUSB* /dev/ttyACM* 2>/dev/null")
    callout(
        doc,
        "设备映射",
        "真机驱动容器运行时再增加 --device=/dev/ttyUSB0:/dev/ttyUSB0；不要给容器 privileged: true。",
        "warning",
    )


def add_day1_4(doc):
    page_break(doc)
    heading(doc, "第 2 章｜Day 1–4 到底做了什么", 1)
    heading(doc, "2.1 七个包的职责", 2)
    table(
        doc,
        ["包", "当前状态", "职责"],
        [
            ("openrobot_description", "已实现", "Xacro、模型、RViz、内部 TF"),
            ("openrobot_gazebo", "已实现", "empty.world、差速插件、关节状态、仿真入口"),
            ("openrobot_bringup", "已实现", "统一选择 sim=true/false"),
            ("openrobot_driver", "占位", "未来串口、运动学、里程计、诊断"),
            ("openrobot_navigation", "占位", "未来 SLAM、AMCL、Nav2"),
            ("openrobot_msgs", "保留", "标准消息不足时才添加自定义接口"),
            ("openrobot_tests", "已实现", "仓库级结构和后续集成测试"),
        ],
        [2400, 1400, 5560],
    )
    heading(doc, "2.2 一条启动命令背后发生什么", 2)
    numbered(doc, "bringup.launch.py 读取 sim 参数；默认 true。")
    numbered(doc, "包含 openrobot_gazebo/launch/sim.launch.py，而不是复制节点定义。")
    numbered(doc, "sim.launch.py 处理 Xacro，生成 robot_description。")
    numbered(doc, "Gazebo 加载 empty.world，spawn_entity 从 robot_description 生成机器人。")
    numbered(doc, "robot_state_publisher 发布机器人内部 TF。")
    numbered(doc, "Gazebo 差速插件订阅 /cmd_vel，控制左右轮，发布 /odom 和 odom→base_footprint。")
    numbered(doc, "Gazebo 关节状态插件发布 /joint_states，robot_state_publisher 据此更新轮子 TF。")
    callout(
        doc,
        "当前缺口",
        "laser_link 只是外形，没有 ray sensor，所以没有 /scan；也没有 map→odom、SLAM、AMCL 或 Nav2。",
        "warning",
    )
    heading(doc, "2.3 当前 Topic 与 TF 所有权", 2)
    table(
        doc,
        ["接口", "当前发布/订阅者", "禁止事项"],
        [
            ("/cmd_vel", "Gazebo 差速插件订阅", "不要同时启动真机驱动消费同一底盘命令"),
            ("/odom", "Gazebo 差速插件发布", "真机模式才改由串口驱动发布"),
            ("odom→base_footprint", "Gazebo 差速插件", "不可再启动第二个发布者"),
            ("/joint_states", "Gazebo 关节插件", "真机模式改由驱动发布"),
            ("内部 TF", "robot_state_publisher", "Gazebo 插件 publish_wheel_tf 必须为 false"),
            ("map→odom", "当前没有", "未来只能由 SLAM Toolbox 或 AMCL 占有"),
        ],
        [2100, 3100, 4160],
    )
    heading(doc, "2.4 亲手复现 Day 1–4", 2)
    command(doc, "C｜容器 Bash", "/workspace", "./scripts/build_ros.sh")
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "source install/setup.bash\n"
        "xacro ros2_ws/src/openrobot_description/urdf/openrobot.urdf.xacro \\\n"
        "  > /tmp/openrobot.urdf\n"
        "check_urdf /tmp/openrobot.urdf",
        "输出 robot name 和完整 link/joint 树，无 ERROR。",
    )
    command(
        doc,
        "C｜终端 1",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 launch openrobot_bringup bringup.launch.py \\\n"
        "  sim:=true use_sim_time:=true use_rviz:=false",
    )
    command(
        doc,
        "C｜终端 2，同一容器",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 topic list\n"
        "ros2 topic echo /odom --once\n"
        "ros2 topic echo /joint_states --once\n"
        "ros2 run tf2_ros tf2_echo odom base_footprint",
        "可看到 /odom、/joint_states；tf2_echo 持续输出变换。",
    )
    command(
        doc,
        "C｜终端 3，同一容器",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 topic pub --once /cmd_vel geometry_msgs/msg/Twist \\\n"
        "  \"{linear: {x: 0.15}, angular: {z: 0.0}}\"",
        "/odom 的 x 应朝正方向变化。若后退，先查 wheel joint axis 和左右映射。",
    )
    checklist(
        doc,
        [
            "我能解释 sim.launch.py 启动的三个节点和一个子 Launch。",
            "我知道 /cmd_vel 的生产者与消费者。",
            "我知道 odom→base_footprint 的唯一发布者。",
            "我能停止 Launch 并重新启动。",
            "我确认当前没有 /scan 和 map→odom。",
        ],
    )


def add_sim_slam(doc):
    page_break(doc)
    heading(doc, "第 3 章｜Day 5–7 激光雷达、办公室世界与 SLAM", 1)
    callout(
        doc,
        "开发方法",
        "每一天都先添加失败测试，再做最小实现，然后运行包级测试和全量测试。"
        "下面给出的文件均是你后续要创建或修改的目标，不代表仓库当前已有。",
        "note",
    )
    day_start(doc, 5, "二维激光雷达插件", "/scan 约 10 Hz，frame_id=laser_link，范围 0.12–8.0 m，RViz 可显示。")
    heading(doc, "3.1.1 修改前检查", 3)
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "rg -n \"ray_sensor|gpu_ray|scan\" ros2_ws/src\n"
        "sed -n '1,220p' ros2_ws/src/openrobot_description/urdf/gazebo.xacro",
        "当前测试明确要求不存在 libgazebo_ros_ray_sensor.so；先修改测试定义新的成功标准。",
    )
    para(doc, "建议选择 CPU `ray`：2D 360 线、10 Hz 对 MVP 足够，兼容无 GPU 的 CI/容器。`gpu_ray` 性能更高，但图形驱动与无头环境更复杂。")
    code(
        doc,
        "加入 gazebo.xacro 的 laser_link 配置（放在 openrobot_gazebo 宏内）：",
        """<gazebo reference="laser_link">
  <sensor name="openrobot_laser" type="ray">
    <always_on>true</always_on>
    <visualize>true</visualize>
    <update_rate>10.0</update_rate>
    <ray>
      <scan>
        <horizontal>
          <samples>360</samples>
          <resolution>1</resolution>
          <min_angle>-3.14159265</min_angle>
          <max_angle>3.14159265</max_angle>
        </horizontal>
      </scan>
      <range>
        <min>0.12</min>
        <max>8.0</max>
        <resolution>0.01</resolution>
      </range>
      <noise>
        <type>gaussian</type>
        <mean>0.0</mean>
        <stddev>0.005</stddev>
      </noise>
    </ray>
    <plugin name="openrobot_laser_plugin"
            filename="libgazebo_ros_ray_sensor.so">
      <ros>
        <remapping>~/out:=scan</remapping>
      </ros>
      <output_type>sensor_msgs/LaserScan</output_type>
      <frame_name>laser_link</frame_name>
    </plugin>
  </sensor>
</gazebo>""",
    )
    callout(
        doc,
        "测试同步修改",
        "test_simulation_assets.py 里原来的“不允许 ray 插件”断言要改为检查插件、scan、laser_link、10 Hz、0.12 和 8.0。不要删除测试。",
        "warning",
    )
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "colcon build --base-paths ros2_ws/src --packages-select \\\n"
        "  openrobot_description openrobot_gazebo --event-handlers console_direct+\n"
        "source install/setup.bash\n"
        "colcon test --base-paths ros2_ws/src --packages-select \\\n"
        "  openrobot_description openrobot_gazebo --event-handlers console_direct+\n"
        "colcon test-result --verbose",
    )
    command(
        doc,
        "C｜仿真运行后新终端",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 topic info /scan -v\n"
        "ros2 topic echo /scan --once\n"
        "ros2 topic hz /scan",
        "LaserScan.header.frame_id 为 laser_link；频率接近 10 Hz；ranges 有有限数值和 inf。",
    )
    callout(
        doc,
        "若 ranges 全是 inf",
        "机器人周围 8 m 内可能没有障碍；先在 Gazebo 放一个箱体。若完全没有 /scan，检查插件库、Xacro 是否安装后重新 source。",
        "note",
    )

    heading(doc, "3.1.2 创建 office_test.world", 3)
    para(doc, "创建 `ros2_ws/src/openrobot_gazebo/worlds/office_test.world`。只使用本地 box，不引用在线 Fuel 模型。")
    code(
        doc,
        "可复用墙体模型片段；每面墙都同时保留 visual 和 collision：",
        """<model name="wall_north">
  <static>true</static>
  <pose>0 3 1 0 0 0</pose>
  <link name="link">
    <collision name="collision">
      <geometry><box><size>8 0.1 2</size></box></geometry>
    </collision>
    <visual name="visual">
      <geometry><box><size>8 0.1 2</size></box></geometry>
      <material><ambient>0.8 0.8 0.8 1</ambient></material>
    </visual>
  </link>
</model>""",
    )
    bullet(doc, "外墙形成约 8 m × 6 m 区域；走廊宽度先设 1.0 m。")
    bullet(doc, "用两段短墙留出 0.85 m 门口，不把门口写成一整面墙。")
    bullet(doc, "加入 0.3–0.5 m 箱体作为固定障碍。")
    bullet(doc, "出生点周围至少留 0.5 m 空间，避免模型一生成就碰撞。")
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "gz sdf -k ros2_ws/src/openrobot_gazebo/worlds/office_test.world\n"
        "ros2 launch openrobot_bringup bringup.launch.py \\\n"
        "  sim:=true use_sim_time:=true use_rviz:=false \\\n"
        "  world:=$(pwd)/ros2_ws/src/openrobot_gazebo/worlds/office_test.world",
        "SDF 检查无错误；机器人生成且 /scan 能看到墙面距离。",
    )

    day_start(doc, 6, "SLAM Toolbox 同步建图", "机器人运动时 /map 更新；TF 出现 map→odom；地图可保存为 YAML+PGM。")
    heading(doc, "3.2.1 新建参数和 Launch", 3)
    para(doc, "创建 `openrobot_navigation/config/slam_params.yaml` 与 `launch/slam.launch.py`，并在 CMakeLists.txt 安装 config/launch。")
    code(
        doc,
        "slam_params.yaml 的稳定起点：",
        """slam_toolbox:
  ros__parameters:
    use_sim_time: true
    mode: mapping
    map_frame: map
    odom_frame: odom
    base_frame: base_footprint
    scan_topic: /scan
    resolution: 0.05
    max_laser_range: 8.0
    minimum_time_interval: 0.2
    transform_publish_period: 0.05
    map_update_interval: 2.0
    minimum_travel_distance: 0.10
    minimum_travel_heading: 0.10
    use_scan_matching: true
    use_scan_barycenter: true""",
    )
    para(doc, "Launch 只启动 `slam_toolbox` 的 `sync_slam_toolbox_node`，不要重复启动 Gazebo。把 `use_sim_time` 和 `params_file` 暴露为参数。")
    command(
        doc,
        "C｜终端 1",
        "/workspace",
        "source install/setup.bash\n./scripts/build_ros.sh\n"
        "ros2 launch openrobot_bringup bringup.launch.py \\\n"
        "  sim:=true use_sim_time:=true use_rviz:=false \\\n"
        "  world:=$(pwd)/ros2_ws/src/openrobot_gazebo/worlds/office_test.world",
    )
    command(
        doc,
        "C｜终端 2",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 launch openrobot_navigation slam.launch.py use_sim_time:=true",
    )
    command(
        doc,
        "C｜终端 3",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 topic hz /scan\n"
        "ros2 topic echo /map --once\n"
        "ros2 run tf2_ros tf2_echo map odom",
        "/map 有 OccupancyGrid；tf2_echo 持续输出 map→odom。此时 map→odom 只能由 SLAM Toolbox 发布。",
    )
    para(doc, "用低速遥控：直线不超过 0.2 m/s，转弯不超过 0.5 rad/s。先绕外墙，再走内部走廊，最后回到起点闭环。")
    command(
        doc,
        "C｜建图完成后",
        "/workspace",
        "mkdir -p maps\n"
        "ros2 run nav2_map_server map_saver_cli -f maps/office_test",
        "生成 maps/office_test.yaml 和 maps/office_test.pgm；重新打开 PGM 应能看到闭合墙体。",
    )
    callout(
        doc,
        "地图重影优先检查",
        "先查 /scan 时间戳、laser_link TF、轮距/轮径和打滑，再调 SLAM 参数。不要用第二个 odom TF 发布者掩盖问题。",
        "risk",
    )

    day_start(doc, 7, "第一周集成", "一条脚本启动仿真，一条脚本启动 SLAM；检查脚本能给出 PASS/WARN/FAIL。")
    para(doc, "创建 `scripts/run_sim.sh`、`scripts/run_slam.sh`、`scripts/check_topics.sh`。统一使用 `set -euo pipefail`，从脚本位置解析仓库根目录，并提供 `--help`。")
    code(
        doc,
        "脚本共同头部：",
        """#!/usr/bin/env bash
set -euo pipefail
script_dir="$(cd -- "$(dirname -- "${BASH_SOURCE[0]}")" && pwd)"
repo_root="$(cd -- "${script_dir}/.." && pwd)"
cd "${repo_root}"
source /opt/ros/humble/setup.bash
if [[ ! -f install/setup.bash ]]; then
  echo "ERROR: run ./scripts/build_ros.sh first" >&2
  exit 1
fi
source install/setup.bash""",
    )
    para(doc, "`run_slam.sh` 只启动 SLAM，不重复启动 Gazebo；`check_topics.sh` 对 /scan、/odom、/map 和四段 TF 分别使用带超时的 ROS CLI。")
    command(doc, "C｜容器 Bash", "/workspace", "bash -n scripts/*.sh\n./scripts/build_ros.sh")
    checklist(
        doc,
        [
            "/scan 平均频率接近 10 Hz，frame_id 正确。",
            "办公室世界不依赖在线模型。",
            "SLAM 运行时只有一个 map→odom 发布者。",
            "地图保存后可以重新打开。",
            "第一周 README 只声明实际完成内容。",
        ],
    )


def add_firmware(doc):
    page_break(doc)
    heading(doc, "第 4 章｜Day 8–14 STM32 电机闭环", 1)
    callout(
        doc,
        "上电红线",
        "首次测试必须架空驱动轮、设置电源限流、低 PWM、随手可断电；先核对 TB6612FNG "
        "连续/峰值电流是否覆盖所购电机堵转电流。不能确认时禁止带负载堵转测试。",
        "risk",
    )
    heading(doc, "4.1 硬件到货当天检查表", 2)
    checklist(
        doc,
        [
            "拍照记录所有包装标签和型号，不凭商品标题猜参数。",
            "万用表确认电源电压与极性。",
            "查电机额定电压、空载电流、堵转电流、减速比和编码器供电。",
            "查 TB6612FNG VM/VCC、电流、STBY 和逻辑电平。",
            "用手转动轮轴，确认没有机械卡死。",
            "编码器 A/B 两相信号先接示波器或逻辑分析仪，不直接进入闭环。",
            "建立参数记录表，未知值写“未测”，不写示例值。",
        ],
    )
    table(
        doc,
        ["参数", "资料值", "实测值", "确认方法"],
        [
            ("电机额定电压", "", "", "电机规格书/卖家书面资料"),
            ("电机堵转电流", "", "", "规格书；不建议自行长时间堵转"),
            ("减速比", "", "", "型号与输出轴转数核对"),
            ("编码器线数", "", "", "规格书+手转计数"),
            ("输出轴 CPR", "", "", "线数×计数边沿×减速比，随后实测一圈"),
            ("轮有效半径", "", "", "负载下滚动多圈测距离"),
            ("轮距", "", "", "左右轮接地点中心距"),
        ],
        [1900, 1800, 1800, 3860],
    )

    day_start(doc, 8, "CubeMX 外设和开环转动", "左右电机可分别低速正反转；复位后 PWM 为 0。")
    numbered(doc, "在 STM32CubeIDE 新建 STM32F407VET6 工程，保留 `.ioc`。")
    numbered(doc, "配置系统时钟；先使用保守频率，确认调试稳定后再优化。")
    numbered(doc, "选择一个定时器的两个 PWM 通道；PWM 频率建议约 20 kHz。")
    numbered(doc, "左右编码器各使用一个支持 Encoder Mode 的定时器。")
    numbered(doc, "配置 USART；初期 115200 8N1，不启用复杂协议。")
    numbered(doc, "配置方向 GPIO、STBY 和状态 LED；上电默认 STBY 关闭、PWM=0。")
    code(
        doc,
        "最小电机接口设计：",
        """typedef enum { MOTOR_COAST, MOTOR_FORWARD, MOTOR_REVERSE, MOTOR_BRAKE } motor_mode_t;

void motor_init(void);
void motor_set_left(motor_mode_t mode, uint16_t pwm);
void motor_set_right(motor_mode_t mode, uint16_t pwm);
void motor_stop_all(void);""",
    )
    callout(
        doc,
        "方向切换",
        "从正转切到反转前先把 PWM 降为 0，等待短暂死区，再切方向；不要在高占空比下直接反向。",
        "warning",
    )
    checklist(doc, ["左轮低 PWM 正转。", "左轮低 PWM 反转。", "右轮低 PWM 正转。", "右轮低 PWM 反转。", "复位、断联时立即停止。"])

    day_start(doc, 9, "编码器计数与方向", "手转左右轮一圈，计数绝对值稳定，前进方向符号一致。")
    para(doc, "先只读计数，不启用 PID。每 10 ms 读取 16/32 位计数器差值，处理回绕。")
    code(
        doc,
        "16 位计数器差值写法：",
        """static int16_t encoder_delta(uint16_t now, uint16_t *last)
{
    int16_t delta = (int16_t)(now - *last);
    *last = now;
    return delta;
}""",
    )
    callout(doc, "符号修正", "优先交换 A/B 相或在一个明确参数处乘以 -1；不要在多个模块重复反号。", "note")

    day_start(doc, 10, "固定周期速度估算", "不同 PWM 下速度趋势合理；单位统一为 mm/s。")
    code(
        doc,
        "速度公式：",
        """distance_per_count_mm = (2 * pi * wheel_radius_mm) / output_shaft_cpr
speed_mm_s = encoder_delta * distance_per_count_mm / sample_period_s""",
    )
    para(doc, "控制周期建议 100 Hz（10 ms）。先记录原始速度，再视噪声加入一阶低通；滤波不能掩盖方向错误或丢计数。")
    table(
        doc,
        ["PWM", "左轮 mm/s", "右轮 mm/s", "电流", "现象"],
        [("10%", "", "", "", ""), ("20%", "", "", "", ""), ("30%", "", "", "", ""), ("40%", "", "", "", "")],
        [1200, 1900, 1900, 1500, 2860],
    )

    day_start(doc, 11, "单电机 PID", "单个电机能稳定跟踪至少三个目标速度，无持续振荡。")
    code(
        doc,
        "PID 状态与一步更新：",
        """typedef struct {
    float kp, ki, kd;
    float integral;
    float prev_error;
    float out_min, out_max;
    float integral_min, integral_max;
} pid_t;

float pid_step(pid_t *p, float target, float measured, float dt)
{
    float error = target - measured;
    p->integral += error * dt;
    if (p->integral > p->integral_max) p->integral = p->integral_max;
    if (p->integral < p->integral_min) p->integral = p->integral_min;
    float derivative = (error - p->prev_error) / dt;
    float out = p->kp * error + p->ki * p->integral + p->kd * derivative;
    p->prev_error = error;
    if (out > p->out_max) out = p->out_max;
    if (out < p->out_min) out = p->out_min;
    return out;
}""",
    )
    numbered(doc, "先确认开环、编码器和采样周期正确。")
    numbered(doc, "Ki=0、Kd=0，从小 Kp 增加到响应快但不持续振荡。")
    numbered(doc, "加入小 Ki 消除稳态误差，同时启用积分限幅。")
    numbered(doc, "只有测量噪声可控且确有需要时才加 Kd。")
    numbered(doc, "每次只改一个参数，记录曲线、超调、稳定时间和稳态误差。")

    day_start(doc, 12, "双电机独立 PID", "相同目标速度下左右轮稳态误差尽量在 5% 内。")
    para(doc, "左右电机使用独立 PID 状态和可独立调节的参数。目标速度先经过斜坡限制，避免一步跳变造成过流。")
    table(
        doc,
        ["试验", "目标", "Kp/Ki/Kd", "超调", "稳定时间", "稳态误差"],
        [("左轮低速", "", "", "", "", ""), ("右轮低速", "", "", "", "", ""), ("双轮中速", "", "", "", "", "")],
        [1300, 1300, 1800, 1300, 1800, 1860],
    )

    day_start(doc, 13, "二进制协议", "电脑可持续发轮速命令并收到反馈；拆包、粘包、坏 CRC 不会失控。")
    table(
        doc,
        ["偏移", "字段", "字节", "说明"],
        [
            ("0", "SOF1", "1", "0xAA"),
            ("1", "SOF2", "1", "0x55"),
            ("2", "Version", "1", "0x01"),
            ("3", "MsgType", "1", "消息类型"),
            ("4", "Sequence", "1", "序列号"),
            ("5", "PayloadLen", "1", "载荷长度"),
            ("6…", "Payload", "N", "小端"),
            ("末尾", "CRC16", "2", "低字节在前"),
        ],
        [1000, 1800, 1000, 5560],
    )
    para(doc, "解析器逐字节寻找 AA 55；长度超限立即丢帧并重新同步；CRC 错误只增加计数，不更新目标速度。")
    code(
        doc,
        "推荐消息：",
        """SET_WHEEL_SPEED:
  int32 left_target_mm_s
  int32 right_target_mm_s
  uint16 command_timeout_ms

MOTOR_FEEDBACK:
  int32 left_encoder_delta
  int32 right_encoder_delta
  int32 left_speed_mm_s
  int32 right_speed_mm_s
  int16 left_pwm
  int16 right_pwm
  uint16 status_flags
  uint32 timestamp_ms""",
    )

    day_start(doc, 14, "通信看门狗与第二周验收", "超过 500 ms 无有效命令自动停车；拔线、坏帧、复位均安全。")
    numbered(doc, "只在收到版本、长度、CRC 都正确的速度命令后刷新 last_valid_command_ms。")
    numbered(doc, "当前时间减 last_valid_command_ms 超过 500 ms，目标速度和积分项清零。")
    numbered(doc, "错误帧、半帧和心跳丢失不得刷新速度命令看门狗。")
    numbered(doc, "上位机重连后先发 STOP/HEARTBEAT，收到新 /cmd_vel 前保持停止。")
    checklist(
        doc,
        [
            "连续运行 10 分钟无异常复位。",
            "拔掉串口后 500 ms 内停车。",
            "发送错误 CRC 不改变电机目标。",
            "固件 README 记录引脚、时钟、周期、协议和安全状态。",
            "保存 PID 曲线和测试数据，不只录视频。",
        ],
    )


def add_driver(doc):
    page_break(doc)
    heading(doc, "第 5 章｜Day 15–21 ROS 2 串口驱动与里程计", 1)
    callout(
        doc,
        "模式互斥",
        "真机驱动发布 odom→base_footprint 时，Gazebo 差速插件必须不运行。统一 Bringup 的 sim 参数就是这个边界。",
        "risk",
    )
    day_start(doc, 15, "C++ 串口封装", "能打开、读写、关闭指定串口；错误日志清晰；配置全部参数化。")
    para(doc, "在 `openrobot_driver` 中增加 `include/openrobot_driver`、`src`、`config` 和测试目录。使用 C++17、rclcpp 和最小 POSIX 串口 API，不引入大型串口框架。")
    code(
        doc,
        "SerialTransport 公开接口：",
        """class SerialTransport {
public:
  SerialTransport();
  ~SerialTransport();
  void open(const std::string & device, int baud_rate);
  void close();
  bool is_open() const noexcept;
  std::size_t read(std::uint8_t * data, std::size_t size);
  void write(const std::uint8_t * data, std::size_t size);
private:
  int fd_{-1};
};""",
    )
    para(doc, "先写伪终端（PTY）测试，再连接 CH340。测试异常包括不存在的设备、断开、短写和 close 后再次 close。")
    command(doc, "W｜Ubuntu-22.04", None, "ls -l /dev/ttyUSB*\ngroups\nusbipd list")
    callout(doc, "权限问题", "容器以设备映射运行；不要用 sudo chmod 777，也不要把整个容器设为 privileged。", "warning")

    day_start(doc, 16, "协议编解码与测试", "CRC、拆包、粘包、垃圾前缀、错误长度、坏版本都通过单元测试。")
    code(
        doc,
        "ProtocolCodec 公开接口：",
        """struct Frame {
  std::uint8_t version;
  std::uint8_t type;
  std::uint8_t sequence;
  std::vector<std::uint8_t> payload;
};

class ProtocolCodec {
public:
  std::vector<std::uint8_t> encode(const Frame & frame) const;
  void feed(const std::uint8_t * data, std::size_t size);
  std::optional<Frame> try_decode();
  static std::uint16_t crc16(const std::uint8_t * data, std::size_t size);
};""",
    )
    table(
        doc,
        ["测试", "输入", "预期"],
        [
            ("正常帧", "完整合法帧", "解出字段与载荷"),
            ("拆包", "每次喂 1–3 字节", "最后才输出一帧"),
            ("粘包", "两帧一次喂入", "连续解出两帧"),
            ("垃圾前缀", "00 FF + 合法帧", "跳过垃圾并同步"),
            ("坏 CRC", "翻转载荷一位", "拒绝并增加错误计数"),
            ("过长载荷", "PayloadLen 超上限", "丢弃并重新找帧头"),
        ],
        [1500, 3800, 4060],
    )

    day_start(doc, 17, "/cmd_vel 与差速逆运动学", "收到 Twist 后生成左右轮 mm/s；限速和命令超时有效。")
    code(
        doc,
        "逆运动学：",
        """left_m_s  = linear_x - angular_z * wheel_separation / 2.0;
right_m_s = linear_x + angular_z * wheel_separation / 2.0;
left_mm_s  = round(left_m_s * 1000.0);
right_mm_s = round(right_m_s * 1000.0);""",
    )
    para(doc, "先对 linear_x 和 angular_z 限幅，再做运动学。参数至少包含 cmd_vel_topic、wheel_radius、wheel_separation、max_linear_velocity、max_angular_velocity 和 cmd_timeout。")
    callout(doc, "ROS 端双保险", "ROS 节点超过 cmd_timeout 没收到 /cmd_vel 时主动发送 STOP；STM32 仍保留自己的 500 ms 看门狗。", "warning")

    day_start(doc, 18, "反馈、正运动学和里程计", "/odom 连续发布，frame_id=odom，child_frame_id=base_footprint。")
    code(
        doc,
        "正运动学与中点积分：",
        """v = (right_m_s + left_m_s) / 2.0;
w = (right_m_s - left_m_s) / wheel_separation;
theta_mid = theta + w * dt / 2.0;
x += v * std::cos(theta_mid) * dt;
y += v * std::sin(theta_mid) * dt;
theta = normalize_angle(theta + w * dt);""",
    )
    para(doc, "dt 使用连续反馈时间戳或单调时钟，必须检查 dt>0 且不异常过大。Odometry pose 和 twist 的 covariance 不要全部留零；按实测逐步调整。")
    table(
        doc,
        ["变换", "仿真模式发布者", "真机模式发布者"],
        [
            ("map→odom", "SLAM/AMCL", "首月真机导航时由 AMCL"),
            ("odom→base_footprint", "Gazebo 差速插件", "openrobot_serial_driver"),
            ("base_footprint→base_link", "robot_state_publisher", "robot_state_publisher"),
        ],
        [2600, 3380, 3380],
    )

    day_start(doc, 19, "/joint_states 与机械标定", "RViz 轮子方向正确；直线距离和原地旋转误差已记录。")
    para(doc, "根据左右累计编码器位置计算轮关节角，发布 left_wheel_joint 与 right_wheel_joint。不要发布 caster 固定关节。")
    numbered(doc, "地面贴胶带，低速直行 1.0 m，测真实距离；修正有效轮半径。")
    numbered(doc, "原地旋转 360°，测真实角度；主要修正 wheel_separation。")
    numbered(doc, "重复至少 5 次，分别记录均值和离散程度。")
    numbered(doc, "一次只改一个参数，改后同时重测直线和旋转。")

    day_start(doc, 20, "自动重连与诊断", "拔插 USB 不崩溃；三秒内尝试恢复；恢复后保持停车。")
    para(doc, "状态机：DISCONNECTED → CONNECTING → CONNECTED → FAULT。读写失败关闭 fd、清接收缓存、停止发送非零目标，按 reconnect_interval 重试。")
    table(
        doc,
        ["诊断字段", "说明"],
        [
            ("connection_state", "connected/disconnected/reconnecting"),
            ("rx_frames / tx_frames", "合法收发帧计数"),
            ("crc_errors", "CRC 错误累计"),
            ("sequence_gaps", "序列号跳变累计"),
            ("reconnect_count", "重连次数"),
            ("last_feedback_age_ms", "反馈新鲜度"),
            ("firmware_status_flags", "固件状态位"),
        ],
        [2800, 6560],
    )

    day_start(doc, 21, "真机链路集成", "/cmd_vel→串口→STM32→电机→编码器→/odom/TF 全链路稳定。")
    command(
        doc,
        "C｜带设备映射的新容器",
        "/workspace",
        'docker run --rm -it \\\n'
        '  --device=/dev/ttyUSB0:/dev/ttyUSB0 \\\n'
        '  -v "$(pwd):/workspace" openrobot-one:humble bash',
    )
    command(
        doc,
        "C｜真机容器",
        "/workspace",
        "source install/setup.bash\n"
        "ros2 launch openrobot_bringup bringup.launch.py \\\n"
        "  sim:=false use_sim_time:=false",
        "hardware.launch.py 应启动真实驱动，而不是当前 Day 4 占位日志；robot_state_publisher 同时运行。",
    )
    checklist(
        doc,
        [
            "轮子架空时先验证方向、停止和超时。",
            "落地后从 0.05 m/s 开始。",
            "停止命令后两轮都可靠停下。",
            "/odom 时间戳使用系统时间，不使用 sim time。",
            "TF 中没有第二个 odom→base_footprint 发布者。",
            "拔线、重连、节点 Ctrl+C 都安全停车。",
        ],
    )


def add_navigation(doc):
    page_break(doc)
    heading(doc, "第 6 章｜Day 22–30 AMCL、Nav2、测试与发布", 1)
    day_start(doc, 22, "加载地图与 AMCL", "保存地图可加载；设置初始位姿后 map→odom 稳定；机器人在 RViz 中与地图对齐。")
    para(doc, "建图和定位不能同时发布 map→odom。启动 AMCL 前停止 SLAM Toolbox。")
    table(
        doc,
        ["组件", "输入", "输出"],
        [
            ("map_server", "YAML+PGM", "/map"),
            ("AMCL", "/scan、/map、odom→base", "map→odom、粒子云、位姿"),
            ("robot_state_publisher", "robot_description、joint_states", "内部 TF"),
        ],
        [2200, 3500, 3660],
    )
    para(doc, "AMCL 参数从 Nav2 Humble 默认配置复制最小相关段，再明确 `base_frame_id=base_footprint`、`odom_frame_id=odom`、`global_frame_id=map`、`scan_topic=scan` 和 `use_sim_time=true`。")
    command(
        doc,
        "C｜定位启动后",
        "/workspace",
        "ros2 lifecycle nodes\n"
        "ros2 topic echo /amcl_pose --once\n"
        "ros2 run tf2_ros tf2_echo map odom",
        "map_server 与 amcl 为 active；amcl_pose 有输出；map→odom 连续存在。",
    )
    callout(doc, "初始位姿", "RViz 的 2D Pose Estimate 要拖出机器人朝向；位置或朝向差太大时 AMCL 可能长期不收敛。", "note")

    day_start(doc, 23, "Nav2 首次启动", "能接收目标点，生成全局路径与局部速度命令；不碰撞。")
    para(doc, "先使用 Nav2 Humble 的标准 bringup 结构，再按本机器人 Frame、Topic 和尺寸做最小修改。不要从空白凭记忆写完整 nav2_params.yaml。")
    table(
        doc,
        ["参数组", "关键检查"],
        [
            ("controller_server", "最大/最小速度、加速度、控制频率"),
            ("planner_server", "规划插件存在，地图可达"),
            ("global_costmap", "global_frame=map，static+obstacle+inflation"),
            ("local_costmap", "global_frame=odom，rolling_window=true"),
            ("footprint", "按底盘真实外形加安全余量，不照搬他人机器人"),
            ("behavior_server", "旋转、后退、等待等恢复行为"),
            ("bt_navigator", "base_frame、odom_topic、目标行为树"),
        ],
        [3100, 6260],
    )
    numbered(doc, "先把最大线速度限制在 0.2 m/s、最大角速度 0.5 rad/s。")
    numbered(doc, "确认静态地图、激光障碍和机器人 footprint 都出现在 costmap。")
    numbered(doc, "发送近距离直线目标，再测试转弯；不要第一次就穿过窄门。")
    numbered(doc, "观察 /cmd_vel、全局路径、局部轨迹和 costmap，而不是只看机器人是否到达。")

    day_start(doc, 24, "前三条基准路线", "直线、90° 转弯和开阔区域路线可重复到达。")
    table(
        doc,
        ["路线", "起点→终点", "限时", "结果", "失败原因"],
        [("R1 直线", "", "", "", ""), ("R2 转弯", "", "", "", ""), ("R3 开阔区", "", "", "", "")],
        [1500, 2800, 1200, 1200, 2660],
    )
    para(doc, "先修定位漂移，再调控制器。若机器人位置在 RViz 跳动，继续调速度只会掩盖根因。")

    day_start(doc, 25, "门口与窄通道", "机器人能通过 0.85–1.0 m 通道，costmap 不把可行通道完全封死。")
    numbered(doc, "实测底盘最大长宽，设置 polygon footprint 和合理 padding。")
    numbered(doc, "显示 global/local costmap，确认 inflation 半径与代价衰减。")
    numbered(doc, "若路径不存在，先看膨胀层；若有路径但摆动，再看控制器与速度。")
    numbered(doc, "一次只调整 footprint、inflation_radius 或 cost_scaling_factor 中一个。")

    day_start(doc, 26, "障碍绕行与恢复行为", "新增障碍后能重规划；无法通过时安全停止或执行有限恢复。")
    para(doc, "分别测试静态障碍、启动后出现的障碍、完全封路。记录规划是否更新、停止距离、恢复次数和最终状态。")
    callout(doc, "安全边界", "恢复旋转前确认 footprint 与周围障碍有余量；真机不要在狭窄区域高速旋转。", "warning")

    day_start(doc, 27, "十条路线统计与参数冻结", "至少执行 10 条标准路线，多轮统计成功率并分类失败。")
    table(
        doc,
        ["编号", "类型", "次数", "成功", "成功率", "主要失败"],
        [
            ("R1–R3", "直线/转弯", "", "", "", ""),
            ("R4–R5", "门口", "", "", "", ""),
            ("R6–R7", "窄通道", "", "", "", ""),
            ("R8–R9", "动态/静态障碍", "", "", "", ""),
            ("R10", "综合路线", "", "", "", ""),
        ],
        [1400, 2000, 1200, 1200, 1500, 2060],
    )
    para(doc, "成功率 = 成功次数 ÷ 总次数 × 100%。失败分类至少区分定位、规划、控制、感知、碰撞、超时和人工终止。")

    day_start(doc, 28, "全新环境复现", "新容器从源码完成 rosdep、build、test 和运行，不依赖个人绝对路径。")
    command(
        doc,
        "W｜Ubuntu-22.04",
        "/mnt/d/OpenRobot-One",
        "docker build --no-cache -f docker/Dockerfile -t openrobot-one:verify .\n"
        'docker run --rm -v "$(pwd):/workspace" openrobot-one:verify \\\n'
        '  bash -lc "./scripts/build_ros.sh"',
        "全量构建测试通过；无手工 apt 补装步骤。",
    )
    command(
        doc,
        "C｜容器 Bash",
        "/workspace",
        "rg -n \"D:\\\\\\\\|/home/|/mnt/d/|ttyUSB0\" ros2_ws/src scripts docker docs",
        "源码中不应存在个人绝对路径；ttyUSB0 只可作为参数默认示例，必须可覆盖。",
    )

    day_start(doc, 29, "README 与作品集证据", "README 只声明实际完成能力，架构图、接线图、曲线和测试表均可追溯。")
    checklist(
        doc,
        [
            "架构图明确仿真轨和真机轨。",
            "Topic/TF 所有权表与代码一致。",
            "硬件接线图标注电源域、共地和 STBY。",
            "PID 曲线包含目标、实测、PWM 和时间轴。",
            "Nav2 测试表包含次数、成功率和失败原因。",
            "常见问题来自实际日志，不虚构。",
            "第三方项目只写参考边界和许可证。",
        ],
    )

    day_start(doc, 30, "最终发布", "Release、v1.0.0 Tag、演示视频和两版讲解稿准备完成。")
    numbered(doc, "从干净环境运行完整构建、测试和关键验收，保存日志。")
    numbered(doc, "录制 2–3 分钟演示：仿真建图/导航、真机闭环、故障安全。")
    numbered(doc, "整理 Release Notes：完成内容、已知限制、运行命令、硬件参数版本。")
    numbered(doc, "确认无密钥、个人串口路径、大型临时文件和生成目录进入提交。")
    numbered(doc, "在你明确决定发布后再提交、推送、打 Tag；本手册不替你自动执行。")
    para(doc, "5 分钟讲解：问题→双轨架构→关键难点→验证数据→结果。15 分钟讲解增加 TF 所有权、协议、PID、里程计和 Nav2 调参过程。")


def add_troubleshooting(doc):
    page_break(doc)
    heading(doc, "第 7 章｜按现象排错", 1)
    table(
        doc,
        ["现象", "第一条检查命令", "优先原因"],
        [
            ("docker 连不上", "docker version", "Docker Desktop 未启动或 WSL 集成关闭"),
            ("ros2 找不到包", "echo $AMENT_PREFIX_PATH", "未 source install/setup.bash"),
            ("colcon 构建失败", "找到第一条 error", "依赖、CMake/package.xml、语法"),
            ("Gazebo 机器人不动", "ros2 topic echo /cmd_vel --once", "Topic、插件、joint 名、轴向"),
            ("没有 /scan", "ros2 topic info /scan -v", "插件未加载、安装未刷新、remap"),
            ("scan 全 inf", "ros2 topic echo /scan --once", "范围内无障碍或碰撞模型缺失"),
            ("TF 断裂", "ros2 run tf2_tools view_frames", "Frame 拼写、缺发布者、sim time"),
            ("地图重影", "ros2 topic hz /scan", "时间戳、里程计、轮距/轮径、打滑"),
            ("串口不存在", "ls -l /dev/ttyUSB*", "USB 未 attach、容器未映射"),
            ("CRC 大量错误", "查看 crc_errors", "波特率、帧边界、字节序、电气噪声"),
            ("电机方向错", "低 PWM 单轮测试", "IN1/IN2、A/B 相、软件符号"),
            ("PID 振荡", "画 target/measured/PWM", "采样、方向、Kp/Ki、饱和"),
            ("里程计转向错", "tf2_echo odom base_footprint", "左右轮交换、符号或轮距"),
            ("AMCL 不收敛", "ros2 topic echo /amcl_pose", "初始位姿、scan/TF、地图"),
            ("Nav2 无路径", "显示 global costmap", "footprint、inflation、地图封闭"),
        ],
        [1900, 3300, 4160],
    )
    heading(doc, "7.1 固定排错顺序", 2)
    numbered(doc, "复现：用最少命令稳定重现，不要同时改多个参数。")
    numbered(doc, "定位第一条失败：构建日志从上往下找第一条 error。")
    numbered(doc, "检查数据是否存在：节点、Topic、频率、消息字段、时间戳。")
    numbered(doc, "检查 TF：连通、方向、唯一发布者、sim time。")
    numbered(doc, "只做一个最小修改，重新运行原验收命令。")
    numbered(doc, "记录原因与证据，再继续下一问题。")
    callout(
        doc,
        "不要这样修",
        "不要增加第二个 TF 发布者、不要 chmod 777、不要 privileged、不要删测试、不要清空整个工作区、不要一次重写多个模块。",
        "risk",
    )

    heading(doc, "7.2 构建失败采集模板", 2)
    code(
        doc,
        "向自己或助手提供以下信息：",
        """环境：Ubuntu 22.04 / ROS 2 Humble / Docker image ID
命令：完整构建命令
第一条 ERROR：原样粘贴
相关文件：package.xml、CMakeLists.txt、报错源文件
git status --short：
已尝试动作：
预期结果：""",
    )
    heading(doc, "7.3 TF 冲突检查", 2)
    command(
        doc,
        "C｜运行中的系统",
        "/workspace",
        "ros2 topic info /tf -v\n"
        "ros2 topic info /tf_static -v\n"
        "ros2 run tf2_ros tf2_echo odom base_footprint\n"
        "ros2 run tf2_tools view_frames",
    )
    para(doc, "如果同一 child frame 在图中出现两个 parent，或 odom→base_footprint 同时来自 Gazebo 与驱动，立即停止其中一个模式。")


def add_appendices(doc):
    page_break(doc)
    heading(doc, "附录 A｜命令速查", 1)
    table(
        doc,
        ["目的", "终端", "命令"],
        [
            ("进入仓库", "W", "cd /mnt/d/OpenRobot-One"),
            ("进入开发容器", "W", "docker compose -f docker/compose.yaml run --rm dev"),
            ("全量构建测试", "C", "./scripts/build_ros.sh"),
            ("加载工作区", "C", "source install/setup.bash"),
            ("启动无头仿真", "C", "ros2 launch openrobot_bringup bringup.launch.py sim:=true use_rviz:=false"),
            ("列 Topic", "C", "ros2 topic list"),
            ("看发布频率", "C", "ros2 topic hz /scan"),
            ("看一次消息", "C", "ros2 topic echo /odom --once"),
            ("查 TF", "C", "ros2 run tf2_ros tf2_echo odom base_footprint"),
            ("保存地图", "C", "ros2 run nav2_map_server map_saver_cli -f maps/office_test"),
            ("检查串口", "W/C", "ls -l /dev/ttyUSB*"),
            ("查看改动", "W", "git status --short --branch && git diff"),
        ],
        [1800, 900, 6660],
        first_col_align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    heading(doc, "附录 B｜Topic 与 TF 最终所有权", 1)
    table(
        doc,
        ["接口", "仿真模式", "真机模式"],
        [
            ("/cmd_vel", "Nav2/遥控发布；Gazebo 订阅", "Nav2/遥控发布；串口驱动订阅"),
            ("/odom", "Gazebo 差速插件", "串口驱动"),
            ("/joint_states", "Gazebo 关节插件", "串口驱动"),
            ("/scan", "Gazebo ray 插件", "实际雷达驱动（后续）"),
            ("map→odom", "SLAM 或 AMCL（二选一）", "AMCL"),
            ("odom→base_footprint", "Gazebo 差速插件", "串口驱动"),
            ("内部 TF", "robot_state_publisher", "robot_state_publisher"),
        ],
        [2500, 3430, 3430],
    )
    heading(doc, "附录 C｜每日收工清单", 1)
    checklist(
        doc,
        [
            "本日成功标准有实际命令或测量证据。",
            "colcon build/test 结果已记录。",
            "未运行的 GUI/硬件测试明确标为未验证。",
            "git diff 只包含本日任务相关文件。",
            "没有删除测试、硬编码个人路径或增加重复 TF。",
            "参数修改记录了修改前、修改后和原因。",
            "电机已停、外部电源已断、USB 状态清楚。",
            "建议提交信息已写好；只有自己确认后才 commit/push。",
        ],
    )
    doc.add_page_break()
    heading(doc, "附录 D｜核心术语", 1)
    table(
        doc,
        ["术语", "通俗解释"],
        [
            ("工作区 overlay", "构建后 source install/setup.bash，让 ROS 找到本项目包"),
            ("Launch", "一次启动多个 ROS 节点并传参的编排文件"),
            ("Topic", "节点之间持续发布/订阅消息的数据通道"),
            ("TF", "坐标系之间随时间变化或固定的空间关系"),
            ("URDF/Xacro", "机器人结构描述；Xacro 可用参数和宏生成 URDF"),
            ("里程计", "根据轮子运动估算机器人相对起点的位置和速度"),
            ("SLAM", "边移动边建图并估计自身位置"),
            ("AMCL", "在已有地图中用粒子滤波定位"),
            ("Costmap", "Nav2 用于规划和避障的代价栅格"),
            ("PID", "用目标与实测误差调节电机输出的闭环控制器"),
            ("CRC16", "检测串口帧在传输中是否损坏的校验值"),
            ("看门狗", "长时间收不到有效命令时强制进入安全状态"),
        ],
        [2400, 6960],
        first_col_align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    callout(
        doc,
        "最后提醒",
        "这份手册给你的是可执行路线，不是跳过理解的粘贴清单。每完成一小步，先解释数据从哪里来、到哪里去，再进入下一步。",
        "success",
    )


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_environment(doc)
    add_day1_4(doc)
    add_sim_slam(doc)
    add_firmware(doc)
    add_driver(doc)
    add_navigation(doc)
    add_troubleshooting(doc)
    add_appendices(doc)
    doc.core_properties.title = "OpenRobot-One Day 5–30 零基础实操手册"
    doc.core_properties.subject = "ROS 2 Humble 双轨移动机器人全流程实操"
    doc.core_properties.author = "OpenRobot-One"
    doc.core_properties.keywords = "ROS 2, Gazebo, STM32, SLAM, Nav2"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
