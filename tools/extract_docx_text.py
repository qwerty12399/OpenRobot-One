from pathlib import Path
import sys

from docx import Document


def extract(path: Path) -> None:
    document = Document(path)
    print(f"# {path.name}")
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)
    for index, table in enumerate(document.tables, start=1):
        print(f"\n## TABLE {index}")
        for row in table.rows:
            print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))


if __name__ == "__main__":
    for raw_path in sys.argv[1:]:
        extract(Path(raw_path))
